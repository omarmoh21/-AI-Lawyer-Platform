"""Generate a formatted .docx file from a filled contract string."""

import io
import logging
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def generate_docx(contract_text: str, title: str) -> bytes:
    """Return a .docx file as bytes from a filled contract text string."""
    doc = Document()

    # ── Page margins ───────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Title ──────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)  # dark blue

    doc.add_paragraph()  # spacer

    # ── Body lines ─────────────────────────────────────────────
    for line in contract_text.strip().splitlines():
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # RTL Arabic

        # Section headers (lines ending with ─ or starting with أولاً etc.)
        is_header = line.startswith(("─", "═")) or (
            len(line) < 60
            and any(line.startswith(p) for p in ["أولاً", "ثانياً", "ثالثاً", "رابعاً",
                                                   "خامساً", "سادساً", "سابعاً", "ثامناً",
                                                   "تاسعاً", "عاشراً"])
        )
        run = para.add_run(line)
        run.font.size = Pt(12)
        if is_header:
            run.bold = True
            run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    # ── Serialize to bytes ─────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    logger.info("docx generated — %d bytes", buf.getbuffer().nbytes)
    return buf.read()
